from http.server import BaseHTTPRequestHandler
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import json
import re
import os

try:
    from upstash_redis import Redis
    redis = Redis(
        url=os.environ.get('UPSTASH_REDIS_REST_URL'),
        token=os.environ.get('UPSTASH_REDIS_REST_TOKEN')
    )
    USE_REDIS = True
except:
    USE_REDIS = False

class handler(BaseHTTPRequestHandler):
    
    def do_GET(self):
        try:
            path_parts = self.path.split('/')
            video_id = path_parts[-1] if len(path_parts) > 0 else None
            
            if not video_id:
                return self.send_error('Missing video_id', 400)
            
            if not USE_REDIS:
                return self.send_error('Cache non disponibile', 503)
            
            cached = redis.get(f'text:{video_id}')
            if not cached:
                return self.send_error('Testo non disponibile', 404)
            
            data = json.loads(cached)
            italian_text = data['italian_text']
            source_lang = data.get('source_lang', 'Sconosciuta')
            
            # Create DOCX
            doc = Document()
            
            # Title
            title = doc.add_heading('Sottotitoli YouTube', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph()
            info = doc.add_paragraph()
            info.add_run('📹 Video ID: ').bold = True
            info.add_run(f'{video_id}\n')
            info.add_run('🌍 Lingua originale: ').bold = True
            info.add_run(f'{source_lang}\n')
            info.add_run('🇮🇹 Tradotto in: ').bold = True
            info.add_run('Italiano\n')
            info.add_run('📅 Data: ').bold = True
            info.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
            
            doc.add_paragraph()
            doc.add_heading('Testo Tradotto', level=1)
            
            # Content
            paragraphs = italian_text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            filename = f'sottotitoli_{video_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(buffer.read())
            
        except Exception as e:
            return self.send_error(str(e), 500)
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
    
    def send_error(self, message, status_code=500):
        self.send_response(status_code)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        response = {'success': False, 'error': message}
        self.wfile.write(json.dumps(response).encode('utf-8'))
